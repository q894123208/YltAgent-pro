"""文档文本提取器：处理 pdf / docx / doc / xlsx / xls / csv 等文本型文件。

这些类型先做确定性的文本/表格抽取，再交给远程多模态模型做医学结构化理解。
"""
from __future__ import annotations

import csv
import io
import logging
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)

DOCX_EXTS = {".docx"}
DOC_EXTS = {".doc"}
XLSX_EXTS = {".xlsx"}
XLS_EXTS = {".xls"}
CSV_EXTS = {".csv"}
PDF_EXTS = {".pdf"}

TEXT_DOC_EXTS = DOCX_EXTS | DOC_EXTS | XLSX_EXTS | XLS_EXTS | CSV_EXTS | PDF_EXTS


def is_text_document(path: Path) -> bool:
    return path.suffix.lower() in TEXT_DOC_EXTS


def extract_docx(path: Path) -> str:
    """python-docx 抽段落和表格。"""
    from docx import Document

    doc = Document(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            parts.append(text)
    for tbl in doc.tables:
        rows: List[str] = []
        for row in tbl.rows:
            cells = [(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("【表格】\n" + "\n".join(rows))
    return "\n".join(parts).strip()


def extract_doc(path: Path) -> str:
    """老版 .doc（OLE 复合文档）尽力抽取；失败给出友好提示。"""
    try:
        import olefile  # noqa: F401
    except ImportError:
        return f"[未能解析 {path.name}：.doc 是旧版二进制格式，建议先用 Word 另存为 .docx 再上传]"
    try:
        import olefile
        ole = olefile.OleFileIO(str(path))
        if ole.exists("WordDocument"):
            stream = ole.openstream("WordDocument").read()
            # 极简提取：扫描可打印的中英文片段
            chunks: List[str] = []
            buf: List[bytes] = []
            for b in stream:
                if 0x20 <= b < 0x7F or b in (0x09, 0x0A, 0x0D) or 0x80 <= b <= 0xFF:
                    buf.append(b)
                else:
                    if len(buf) > 6:
                        try:
                            chunks.append(bytes(buf).decode("gbk", errors="ignore"))
                        except Exception:
                            pass
                    buf = []
            ole.close()
            text = "\n".join([c.strip() for c in chunks if c.strip()])
            if text:
                return text
        ole.close()
    except Exception as exc:
        logger.warning("extract_doc failed for %s: %s", path, exc)
    return f"[未能解析 {path.name}：建议先用 Word 另存为 .docx 再上传]"


def extract_xlsx(path: Path) -> str:
    """openpyxl 抽每个 sheet 的所有非空单元格，按行拼成 markdown 风格。"""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True, read_only=True)
    parts: List[str] = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"【工作表：{sheet_name}】")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts).strip()


def extract_xls(path: Path) -> str:
    """老版 .xls 用 xlrd 1.x（注：xlrd ≥ 2.0 已不支持 xlsx）。"""
    try:
        import xlrd
    except ImportError:
        return f"[未能解析 {path.name}：缺少 xlrd 依赖，建议先另存为 .xlsx]"
    try:
        book = xlrd.open_workbook(str(path))
    except Exception as exc:
        return f"[未能解析 {path.name}：{type(exc).__name__}]"
    parts: List[str] = []
    for sheet in book.sheets():
        parts.append(f"【工作表：{sheet.name}】")
        for row_idx in range(sheet.nrows):
            row = sheet.row_values(row_idx)
            cells = ["" if v in (None, "") else str(v) for v in row]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_csv(path: Path) -> str:
    """CSV：尝试 utf-8 / gbk / utf-8-sig，按行拼接。"""
    raw_bytes = path.read_bytes()
    text = ""
    for encoding in ("utf-8-sig", "utf-8", "gbk", "gb18030", "latin-1"):
        try:
            text = raw_bytes.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if not text:
        return f"[未能识别 {path.name} 编码]"
    parts: List[str] = []
    reader = csv.reader(io.StringIO(text))
    for row in reader:
        cells = [(c or "").strip() for c in row]
        if any(cells):
            parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_pdf(path: Path) -> str:
    """PDF 文本层探测：能提取文字就返回；扫描型页留给 VLM 图片解析。"""
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("缺少依赖 PyMuPDF，请 pip install pymupdf") from exc

    parts: List[str] = []
    with fitz.open(str(path)) as doc:
        for page_idx, page in enumerate(doc, start=1):
            text = (page.get_text("text") or "").strip()
            if text:
                parts.append(f"【PDF 第 {page_idx} 页｜文本层】\n{text}")
            else:
                parts.append(f"【PDF 第 {page_idx} 页｜未提取到文字】")
    return "\n\n".join(parts).strip()


def extract_text(path: Path) -> str:
    """根据扩展名分发到具体抽取函数。"""
    suffix = path.suffix.lower()
    if suffix in DOCX_EXTS:
        return extract_docx(path)
    if suffix in DOC_EXTS:
        return extract_doc(path)
    if suffix in XLSX_EXTS:
        return extract_xlsx(path)
    if suffix in XLS_EXTS:
        return extract_xls(path)
    if suffix in CSV_EXTS:
        return extract_csv(path)
    if suffix in PDF_EXTS:
        return extract_pdf(path)
    return ""
